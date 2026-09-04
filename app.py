import streamlit as st
from langchain_ollama import ChatOllama
from langchain_core.prompts import PromptTemplate
from pypdf import PdfReader
from docx import Document
from PIL import Image
import pypdfium2 as pdfium
import numpy as np
import os
import sys
import io
import time
import json
import urllib.request

st.set_page_config(
    page_title="Trợ Lý AI: Đánh Giá Tài Liệu",
    page_icon="📑",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Đường dẫn thư mục Startup của Windows để mở cùng Windows
STARTUP_DIR = os.path.join(os.environ.get("APPDATA", ""), r"Microsoft\Windows\Start Menu\Programs\Startup")
STARTUP_BAT = os.path.join(STARTUP_DIR, "AI_Document_Assistant.bat")

def is_autostart_enabled():
    return os.path.exists(STARTUP_BAT)

def toggle_autostart(enable: bool):
    try:
        if enable:
            python_exe = sys.executable
            app_dir = os.path.abspath(os.path.dirname(__file__))
            bat_content = f"""@echo off
cd /d "{app_dir}"
timeout /t 5 /nobreak >nul
start "" "http://localhost:8501"
"{python_exe}" -m streamlit run app.py --server.headless=true
"""
            with open(STARTUP_BAT, "w", encoding="utf-8") as f:
                f.write(bat_content)
            return True
        else:
            if os.path.exists(STARTUP_BAT):
                os.remove(STARTUP_BAT)
            return True
    except Exception as e:
        st.sidebar.error(f"Lỗi cài đặt khởi động cùng Windows: {e}")
        return False

# Lấy danh sách mô hình đang có sẵn trong Ollama
def get_installed_models():
    try:
        req = urllib.request.urlopen("http://localhost:11434/api/tags", timeout=3)
        data = json.loads(req.read().decode("utf-8"))
        models = [m["name"].split(":")[0] + ":" + m["name"].split(":")[1] if ":" in m["name"] else m["name"] for m in data.get("models", [])]
        models = list(dict.fromkeys(models))
        if models:
            # Sắp xếp ưu tiên mô hình nhẹ trước: 1.5b -> 3b -> 7b
            priority = {"qwen2.5:1.5b": 0, "qwen2.5:3b": 1, "qwen2.5:7b": 2}
            models.sort(key=lambda x: priority.get(x, 99))
            return models
    except Exception:
        pass
    return ["qwen2.5:1.5b", "qwen2.5:3b", "qwen2.5:7b"]

# Khởi tạo công cụ OCR (lưu cache trong bộ nhớ để không tải lại nhiều lần)
@st.cache_resource
def get_ocr_engine():
    from rapidocr_onnxruntime import RapidOCR
    return RapidOCR()

# Trích xuất nội dung văn bản (Hỗ trợ tài liệu số và tài liệu SCAN ảnh qua OCR)
def extract_text_from_file(uploaded_file):
    text = ""
    is_scanned = False
    filename = uploaded_file.name.lower()
    file_bytes = uploaded_file.read()
    uploaded_file.seek(0)

    if filename.endswith(".pdf"):
        # 1. Thử trích xuất văn bản số thông thường trước
        try:
            reader = PdfReader(io.BytesIO(file_bytes))
            total_pages = len(reader.pages)
            for page in reader.pages:
                page_text = page.extract_text() or ""
                text += page_text + "\n"
        except Exception:
            total_pages = 0

        # 2. Kiểm tra nếu là PDF scan (ít hoặc không có chữ số)
        stripped_len = len(text.strip())
        if total_pages > 0 and (stripped_len < 50 or (stripped_len / total_pages < 30)):
            is_scanned = True
            ocr_text = []
            try:
                ocr = get_ocr_engine()
                doc = pdfium.PdfDocument(file_bytes)
                # Giới hạn xử lý tối đa 10 trang đầu để đảm bảo tốc độ
                max_pages = min(len(doc), 10)
                for i in range(max_pages):
                    page = doc[i]
                    bitmap = page.render(scale=1.5)
                    np_img = bitmap.to_numpy()
                    res, _ = ocr(np_img)
                    if res:
                        lines = [box[1] for box in res if box[1].strip()]
                        if lines:
                            ocr_text.append(f"[Trang {i+1}]\n" + "\n".join(lines))
                if ocr_text:
                    text = "\n\n".join(ocr_text)
            except Exception as e:
                if not text:
                    text = f"[Lỗi OCR: {e}]"

    elif filename.endswith((".png", ".jpg", ".jpeg")):
        is_scanned = True
        try:
            ocr = get_ocr_engine()
            img = Image.open(io.BytesIO(file_bytes)).convert("RGB")
            res, _ = ocr(np.array(img))
            if res:
                lines = [box[1] for box in res if box[1].strip()]
                text = "\n".join(lines)
        except Exception as e:
            text = f"[Lỗi OCR hình ảnh: {e}]"

    elif filename.endswith(".docx"):
        doc = Document(io.BytesIO(file_bytes))
        for p in doc.paragraphs:
            text += p.text + "\n"

    elif filename.endswith(".txt"):
        text = file_bytes.decode("utf-8", errors="ignore")

    return text.strip(), is_scanned

# ==================== GIAO DIỆN THANH BÊN (SIDEBAR) ====================
with st.sidebar:
    st.header("⚙️ Cấu Hình Mô Hình AI")
    
    available_models = get_installed_models()
    
    default_index = 0
    if "qwen2.5:1.5b" in available_models:
        default_index = available_models.index("qwen2.5:1.5b")
    elif "qwen2.5:3b" in available_models:
        default_index = available_models.index("qwen2.5:3b")

    selected_model = st.selectbox(
        "Chọn mô hình AI:",
        options=available_models,
        index=default_index,
        help="Mô hình 1.5B/3B chạy cực nhanh trên CPU i3. Mô hình 7B nặng hơn nhiều."
    )

    if "1.5b" in selected_model:
        st.success("🚀 **Siêu tốc**: Mô hình cực nhẹ (~986MB), phản hồi nhanh nhất trên CPU!")
    elif "3b" in selected_model:
        st.success("⚡ **Tốc độ cao & Cân bằng**: Phù hợp cho CPU i3 và RAM 8GB.")
    else:
        st.warning("⚠️ Mô hình 7B khá nặng khi chạy trên CPU i3, tốc độ sinh từ sẽ chậm hơn.")

    st.markdown("---")
    st.subheader("⚡ Tối Ưu Tốc Độ Xử Lý")
    
    analysis_mode = st.radio(
        "Chế độ phản hồi:",
        options=["⚡ Siêu tốc (Súc tích, gạch đầu dòng ngắn gọn)", "📋 Tiêu chuẩn (Chi tiết, phân tích sâu)"],
        index=0,
        help="Chế độ 'Siêu tốc' giảm lượng từ cần sinh ra, giúp hoàn thành nhanh hơn gấp đôi!"
    )

    context_limit = st.slider(
        "Giới hạn ký tự tài liệu phân tích:",
        min_value=2000,
        max_value=12000,
        value=5000,
        step=500,
        help="4.000 - 5.000 ký tự (~1.000 - 1.500 từ) là độ dài lý tưởng để CPU không bị trễ."
    )
    
    temperature = st.slider(
        "Độ sáng tạo (Temperature):",
        min_value=0.0,
        max_value=1.0,
        value=0.2,
        step=0.05,
        help="0.2 giúp câu trả lời súc tích, logic và bám sát tài liệu."
    )

    st.markdown("---")
    st.subheader("🖥️ Tính Năng Hệ Thống")
    autostart_current = is_autostart_enabled()
    autostart_toggle = st.checkbox(
        "🚀 Mở cùng Windows (Khởi động tự động)",
        value=autostart_current,
        help="Tự động khởi động ứng dụng mỗi khi bạn bật máy tính Windows."
    )
    
    if autostart_toggle != autostart_current:
        if toggle_autostart(autostart_toggle):
            st.toast("✅ Đã cập nhật cài đặt khởi động cùng Windows!", icon="🚀")
            st.rerun()

    st.caption("Tip: Bạn cũng có thể nhấp đúp vào file `start_app.bat` trong thư mục để khởi động nhanh.")


# ==================== GIAO DIỆN CHÍNH ====================
st.title("📑 Trợ Lý AI: Đánh Giá & Nhận Xét Báo Cáo")
st.markdown("*Hỗ trợ văn bản điện tử và **bản SCAN / hình ảnh qua OCR Offline** bảo mật.*")

uploaded_file = st.file_uploader(
    "Chọn tài liệu cần đánh giá (Hỗ trợ PDF thông thường, PDF SCAN, DOCX, TXT, PNG, JPG)",
    type=["pdf", "docx", "txt", "png", "jpg", "jpeg"]
)

if uploaded_file:
    # Đọc và tính toán sơ bộ tài liệu (có tích hợp OCR nếu là tài liệu scan)
    if "cached_file_name" not in st.session_state or st.session_state["cached_file_name"] != uploaded_file.name:
        with st.spinner("🔍 Đang đọc và nhận dạng văn bản (tự động nhận diện OCR nếu là bản Scan)..."):
            extracted_text, is_scanned = extract_text_from_file(uploaded_file)
            st.session_state["cached_text"] = extracted_text
            st.session_state["cached_file_name"] = uploaded_file.name
            st.session_state["is_scanned"] = is_scanned
            st.session_state["report_content"] = None

    raw_text = st.session_state.get("cached_text", "")
    is_scanned = st.session_state.get("is_scanned", False)
    char_count = len(raw_text)
    word_count = len(raw_text.split())

    col1, col2, col3 = st.columns(3)
    col1.metric("Tên tài liệu", uploaded_file.name)
    col2.metric("Số từ nhận dạng", f"{word_count:,} từ")
    col3.metric("Số ký tự", f"{char_count:,} ký tự")

    if is_scanned:
        st.info("🔍 **Đã phát hiện tài liệu SCAN (ảnh chụp)**: Hệ thống đã sử dụng bộ máy OCR Offline để tự động đọc toàn bộ chữ từ hình ảnh.")

    if not raw_text:
        st.error("⚠️ Không thể trích xuất được chữ từ tài liệu này. Vui lòng kiểm tra lại chất lượng bản scan hoặc chọn file khác.")
    else:
        with st.expander("👁️ Xem trước nội dung văn bản trích xuất được (để kiểm tra trước khi AI phân tích):"):
            st.text_area("Nội dung nhận dạng:", value=raw_text[:3000] + ("\n... [Đã rút gọn hiển thị]" if len(raw_text) > 3000 else ""), height=150, disabled=True)

        if char_count > context_limit:
            st.caption(f"ℹ️ Tài liệu dài {char_count:,} ký tự. Hệ thống sẽ trích xuất {context_limit:,} ký tự đầu tiên để phân tích nhanh nhất.")

        start_analyze = st.button("🚀 Bắt đầu phân tích & Đánh giá", type="primary", use_container_width=True)

        if start_analyze:
            doc_context = raw_text[:context_limit]

            # Tùy chỉnh phong cách prompt theo chế độ
            if "Siêu tốc" in analysis_mode:
                style_instruction = """- Trình bày dạng các gạch đầu dòng súc tích, ngắn gọn, đi thẳng vào vấn đề.
- Không viết đoạn văn mở đầu hay kết thúc rườm rà.
- Trực diện, mỗi mục từ 2 - 4 ý cốt lõi nhất."""
            else:
                style_instruction = """- Phân tích chi tiết, đầy đủ luận cứ và dẫn chứng cụ thể từ tài liệu.
- Viết văn phong trang trọng, chuẩn mực."""

            prompt_template = PromptTemplate(
                input_variables=["doc", "style"],
                template="""Bạn là một chuyên gia phân tích và đánh giá tài liệu chuyên nghiệp.
Dưới đây là nội dung tài liệu được cung cấp:

---
{doc}
---

Hãy xuất ra một bản **BÁO CÁO ĐÁNH GIÁ VÀ NHẬN XÉT** theo cấu trúc chuẩn sau:
1. **TỔNG QUAN TÀI LIỆU**: Chủ đề chính, mục đích và phạm vi nội dung.
2. **TÓM TẮT CÁC Ý CHÍNH**: Các luận điểm/nội dung quan trọng nhất.
3. **ƯU ĐIỂM & ĐIỂM NỔI BẬT**: Những phần viết tốt, dữ liệu thuyết phục hoặc ý tưởng sáng tạo.
4. **HẠN CHẾ & THIẾU SÓT**: Lỗi lập luận, thiếu chứng cứ, định dạng chưa chuẩn hoặc điểm mơ hồ.
5. **ĐÁNH GIÁ & KIẾN NGHỊ CẢI THIỆN**: Hướng khắc phục cụ thể để hoàn thiện tài liệu.

Yêu cầu phong cách trình bày:
{style}

Định dạng văn bản bằng Markdown rõ ràng, dễ đọc."""
            )

            # Cấu hình tối ưu 8 luồng CPU và bộ nhớ đệm cho Ollama
            llm = ChatOllama(
                model=selected_model,
                temperature=temperature,
                num_thread=8,
                num_ctx=4096
            )
            chain = prompt_template | llm

            # ================= TRẠNG THÁI SUY NGHĨ (THINKING & LOAD STATUS) =================
            with st.status("🧠 AI đang phân tích và hình thành luồng suy nghĩ...", expanded=True) as status:
                st.write("📖 **Bước 1: Khởi tạo ngữ cảnh** — Đã nạp dữ liệu văn bản vào bộ đệm.")
                st.write(f"⚙️ **Bước 2: Tối ưu phần cứng** — Kích hoạt 8 luồng CPU cho mô hình `{selected_model}`.")
                st.write("🤔 **Bước 3: Quá trình suy luận** — Đang phân tích luận điểm, đối chiếu ưu/nhược điểm và xây dựng bố cục báo cáo...")
                
                start_time = time.time()
                
                def stream_generator():
                    for chunk in chain.stream({"doc": doc_context, "style": style_instruction}):
                        if hasattr(chunk, "content"):
                            yield chunk.content
                        else:
                            yield str(chunk)

                # Cập nhật trạng thái khi bắt đầu sinh kết quả
                status.update(label=f"✍️ Đang xuất báo cáo trực tiếp bằng mô hình {selected_model}...", state="running", expanded=False)

            st.markdown(f"### 📊 Kết quả đánh giá chi tiết (Mô hình: `{selected_model}`):")

            # Khối hiển thị chữ chạy trực tiếp (Streaming)
            full_content = st.write_stream(stream_generator)
            elapsed_time = round(time.time() - start_time, 1)

            # Đánh dấu hoàn thành trong status
            st.toast(f"✅ Hoàn thành phân tích trong {elapsed_time}s!", icon="⚡")

            # Lưu vào session_state
            st.session_state["report_content"] = full_content
            st.session_state["report_file_name"] = f"Danh_Gia_{uploaded_file.name.rsplit('.', 1)[0]}.md"
            st.session_state["elapsed_time"] = elapsed_time
            st.session_state["used_model"] = selected_model

        # Hiển thị báo cáo và nút tải nếu đã có kết quả
        if st.session_state.get("report_content") and not start_analyze:
            st.markdown(f"### 📊 Kết quả đánh giá chi tiết (Mô hình: `{st.session_state.get('used_model')}`):")
            st.markdown(st.session_state["report_content"])

        if st.session_state.get("report_content"):
            st.success(f"⚡ Đã hoàn thành đánh giá trong **{st.session_state.get('elapsed_time')} giây** với mô hình **{st.session_state.get('used_model')}**!")
            
            st.download_button(
                label="📥 Tải báo cáo đánh giá (.md)",
                data=st.session_state["report_content"],
                file_name=st.session_state.get("report_file_name", "Bao_Cao_Danh_Gia.md"),
                mime="text/markdown",
                use_container_width=True
            )